import requests
import base64
import json
import google.generativeai as genai
import docx
import io
import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
import time
import re
from datetime import datetime
from config import Config
import os
import logging

logger = logging.getLogger(__name__)


class ObservationExtractor:
    def __init__(self):
        self.ocr_api_key = Config.OCR_API_KEY
        self.groq_api_key = Config.GROQ_API_KEY
        self.gemini_api_key = Config.GOOGLE_API_KEY
        genai.configure(api_key=self.gemini_api_key)

    def get_pronouns(self, gender):
        """Get appropriate pronouns based on gender"""
        if gender == "male":
            return {"subject": "he", "object": "him", "possessive": "his"}
        elif gender == "female":
            return {"subject": "she", "object": "her", "possessive": "her"}
        else:
            return {"subject": "they", "object": "them", "possessive": "their"}

    def image_to_base64(self, image_file):
        """Convert image file to base64 string"""
        return base64.b64encode(image_file.read()).decode("utf-8")

    def extract_text_with_ocr(self, image_file):
        """Extract text from image using OCR.space API"""
        try:
            # Get file extension
            file_type = image_file.filename.split(".")[-1].lower()
            if file_type == "jpeg":
                file_type = "jpg"

            # Convert image to base64
            image_file.seek(0)  # Reset file pointer
            base64_image = self.image_to_base64(image_file)
            base64_image_with_prefix = f"data:image/{file_type};base64,{base64_image}"

            # Prepare request payload
            payload = {
                "apikey": self.ocr_api_key,
                "language": "eng",
                "isOverlayRequired": False,
                "iscreatesearchablepdf": False,
                "issearchablepdfhidetextlayer": False,
                "OCREngine": 2,
                "detectOrientation": True,
                "scale": True,
                "base64Image": base64_image_with_prefix,
            }

            # Send request to OCR API
            response = requests.post(
                "https://api.ocr.space/parse/image",
                data=payload,
                headers={"apikey": self.ocr_api_key},
            )

            response.raise_for_status()
            data = response.json()

            # Process response
            if not data.get("ParsedResults") or len(data["ParsedResults"]) == 0:
                error_msg = data.get("ErrorMessage", "No parsed results returned")
                raise Exception(f"OCR Error: {error_msg}")

            parsed_result = data["ParsedResults"][0]
            if parsed_result.get("ErrorMessage"):
                raise Exception(f"OCR Error: {parsed_result['ErrorMessage']}")

            extracted_text = parsed_result["ParsedText"]

            if not extracted_text or not extracted_text.strip():
                raise Exception("No text was detected in the image")

            return extracted_text

        except Exception as e:
            raise Exception(f"OCR Error: {str(e)}")

    def process_with_groq(self, extracted_text):
        """Process extracted text with Groq AI"""
        try:
            system_prompt = """You are an AI assistant for a learning observation system. Extract and structure information from the provided observation sheet text.

CRITICAL INSTRUCTIONS FOR NAME HANDLING:
- Do NOT use any names that appear in the observation text or audio transcription
- For the "studentName" field, ONLY use names that are explicitly provided by the system/database
- If no database name is provided, use "Student" as the default
- NEVER assume gender - always refer to the student as "the student" or "Student" in descriptions
- Do not use gender-specific pronouns (he/his, she/her) in any part of the response

The observation sheets typically have the following structure:
- Title (usually "The Observer")
- Student information (Name, Roll Number/ID) - IGNORE names from this section
- Date and Time information
- Core Observation Section with time slots
- Teaching content for each time slot
- Learning details (what was learned, tools used, etc.)

Format your response as JSON with the following structure:
{
  "studentName": "Use ONLY the database-provided name, never from observation text",
  "studentId": "Student ID or Roll Number",
  "className": "Class name or subject being taught",
  "date": "Date of observation",
  "observations": "Detailed description of what was learned - refer to 'the student' not by name",
  "strengths": ["List of strengths observed - use 'the student' in descriptions"],
  "areasOfDevelopment": ["List of areas where the student needs improvement - use 'the student'"],
  "recommendations": ["List of recommended actions - refer to 'the student'"],
  "themeOfDay": "Main theme or topic of the day",
  "curiositySeed": "Something that sparked the child's interest"
}

For observations, provide full detailed descriptions like:
"The student learned how to make maggi from their parent through in-person mode, including all steps from boiling water to adding spices"

IMPORTANT: Never use gender-specific language or names from the observation text. Always refer to 'the student' in descriptions."""

            # Send request to Groq API
            response = requests.post(
                "https://api.groq.com/openai/v1/chat/completions",
                headers={
                    "Authorization": f"Bearer {self.groq_api_key}",
                    "Content-Type": "application/json",
                },
                json={
                    "model": "llama-3.3-70b-versatile",
                    "messages": [
                        {"role": "system", "content": system_prompt},
                        {
                            "role": "user",
                            "content": f"Extract and structure the following text from an observation sheet: {extracted_text}",
                        },
                    ],
                    "temperature": 0.2,
                    "response_format": {"type": "json_object"},
                },
            )

            response.raise_for_status()
            data = response.json()

            # Extract the JSON content
            ai_response = data["choices"][0]["message"]["content"]
            return json.loads(ai_response)

        except Exception as e:
            raise Exception(f"Groq API Error: {str(e)}")

    def transcribe_with_assemblyai(self, audio_file, language_code="en"):
        """Transcribe audio using AssemblyAI API with English, Hindi, Marathi, or Punjabi."""
        if not Config.ASSEMBLYAI_API_KEY:
            return "Error: AssemblyAI API key is not configured."

        headers = {
            "authorization": Config.ASSEMBLYAI_API_KEY,
            "content-type": "application/json",
        }

        try:
            # Upload the audio file
            audio_file.seek(0)
            upload_response = requests.post(
                "https://api.assemblyai.com/v2/upload",
                headers={"authorization": Config.ASSEMBLYAI_API_KEY},
                data=audio_file.read(),
            )

            if upload_response.status_code != 200:
                return f"Error uploading audio: {upload_response.text}"

            upload_url = upload_response.json()["upload_url"]

            # Prepare transcription request
            transcript_request = {
                "audio_url": upload_url,
                "language_code": language_code,  # 'en', 'hi', 'mr', or 'pa'
            }

            transcript_response = requests.post(
                "https://api.assemblyai.com/v2/transcript",
                json=transcript_request,
                headers=headers,
            )

            if transcript_response.status_code != 200:
                return f"Error requesting transcription: {transcript_response.text}"

            transcript_id = transcript_response.json()["id"]

            status = "processing"
            while status != "completed" and status != "error":
                polling_response = requests.get(
                    f"https://api.assemblyai.com/v2/transcript/{transcript_id}",
                    headers=headers,
                )

                if polling_response.status_code != 200:
                    return (
                        f"Error checking transcription status: {polling_response.text}"
                    )

                polling_data = polling_response.json()
                status = polling_data["status"]

                if status == "completed":
                    return polling_data["text"]
                elif status == "error":
                    return f"Transcription error: {polling_data.get('error', 'Unknown error')}"

                time.sleep(2)

            return "Error: Transcription timed out or failed."
        except Exception as e:
            return f"Error during transcription: {str(e)}"

    def generate_conversational_transcript(self, raw_text):
        """Convert raw observations/transcript into conversational format using Gemini API"""
        try:
            prompt = f"""
            Convert the following observation text into a conversational format between an Observer and a Child. 

CRITICAL INSTRUCTIONS:
- NEVER use any names that appear in the raw text or audio
- Always refer to the child as "Child" in the dialogue labels
- Do not assume gender - avoid using he/his, she/her pronouns
- Use gender-neutral language throughout the conversation

Format it as a natural dialogue where:
            - Observer speaks first with questions, instructions, or observations
            - Child responds naturally based on the context
            - Use "Observer:" and "Child:" labels for each speaker
            - Make it realistic and age-appropriate
            - If the text is already conversational, just format it properly
            - If it's narrative observations, convert them into likely dialogue
            - Create a natural flow of conversation that would lead to the observations described
            - Include educational moments and learning interactions
            - NEVER use names from the original text - always use "Child:" as the label
            - Avoid gender assumptions in the dialogue content

            Original observation text:
            {raw_text}

            Please format as a realistic conversation:
            Observer: [what the observer might have said or asked]
            Child: [how the child might have responded based on the context]
            Observer: [follow-up from observer]
            Child: [child's response]

            Keep it natural, educational, and age-appropriate. Make sure the conversation flows logically and would realistically result in the observations described in the original text. Remember to never use names from the original text and avoid gender-specific language.
            """

            # Use the same Gemini API pattern as your existing methods
            model = genai.GenerativeModel("gemini-2.0-flash")
            config = genai.types.GenerationConfig(temperature=0.2)
            response = model.generate_content(
                [{"role": "user", "parts": [{"text": prompt}]}],
                generation_config=config,
            )

            if response and response.text:
                return response.text.strip()
            else:
                return self._basic_transcript_formatting(raw_text)

        except Exception as e:
            # Fallback to basic formatting if API fails
            return self._basic_transcript_formatting(raw_text)

    def _basic_transcript_formatting(self, raw_text):
        """Basic fallback transcript formatting"""
        lines = raw_text.split("\n")
        formatted_lines = []

        for i, line in enumerate(lines):
            line = line.strip()
            if line:
                if i % 2 == 0:
                    formatted_lines.append(f"Observer: {line}")
                else:
                    formatted_lines.append(f"Child: {line}")

        if not formatted_lines:
            # If no lines, create a basic conversation
            formatted_lines = [
                "Observer: Can you tell me about what you learned today?",
                f"Child: {raw_text[:100]}..."
                if len(raw_text) > 100
                else f"Child: {raw_text}",
                "Observer: That's wonderful! Can you tell me more about it?",
                "Child: Yes, I enjoyed learning about this topic.",
            ]

        return "\n".join(formatted_lines)

    def generate_report_from_text(self, text_content, user_info):
        """Generate a structured report from text using Google Gemini"""
        # Get child information with gender
        from models.database import get_child_by_id

        # Ensure we have a valid child_id and student_name
        child_id = user_info.get("child_id")
        student_name = user_info.get("student_name", "Student")

        # If no child_id, try to get child by name
        if not child_id and student_name != "Student":
            from models.database import get_supabase_client

            supabase = get_supabase_client()
            child_data = (
                supabase.table("children")
                .select("id, name, gender")
                .eq("name", student_name)
                .execute()
                .data
            )
            if child_data:
                child_id = child_data[0]["id"]
                child = child_data[0]
            else:
                child = None
        else:
            child = get_child_by_id(child_id) if child_id else None

        # Get pronouns based on gender
        pronouns = (
            self.get_pronouns(child["gender"])
            if child and child.get("gender")
            else {"subject": "they", "object": "them", "possessive": "their"}
        )

        # Ensure we have a valid student name
        if not student_name or student_name == "Student":
            student_name = child["name"] if child and child.get("name") else "Student"

        prompt = f"""
        You are an educational observer tasked with generating a comprehensive and accurate Daily Insights based on the following observational notes from a student session. Pay special attention to any achievements, learning moments, and areas for growth. The report should be structured, insightful, and easy to understand for parents. Add postives and negatives based on the text content provided.

        CRITICAL INSTRUCTIONS FOR NAME AND GENDER USAGE:
        - NEVER extract or use any name from the audio transcription or text content
        - ALWAYS use the exact name provided: {student_name}
        - Use these pronouns for the student throughout the report: subject = {pronouns["subject"]}, object = {pronouns["object"]}, possessive = {pronouns["possessive"]}
        - When referring to the student, use "{student_name}" or the appropriate pronouns ({pronouns["subject"]}/{pronouns["object"]}/{pronouns["possessive"]})
        - Make sure the report is grammatically correct and adheres to proper English syntax and semantics.
        Please carefully analyze the given text and complete the report using the exact format, emojis, section titles, and scoring rubrics as described below. The student should be referred to consistently using their provided name "{student_name}" and the appropriate pronouns - never use names from the audio/text content.

        📌 Important Instructions for the Report:
        - Follow the format exactly as shown below.
        - Make reasonable inferences for items not explicitly stated in the text.
        - Ensure that the final Overall Growth Score and category (🟢/💚/⚠️/📈) accurately reflects the number of active areas, according to:
        🟢 Excellent (7/7 areas) – Clear growth with strong evidence
        💚 Good (5-6 areas) – Solid engagement with positive trends
        ⚠️ Fair (3-4 areas) – Some engagement, needs encouragement
        📈 Needs Work (1-2 areas) – Area not activated or underperforming today
        - Include the new Communication Skills & Thought Clarity section.
        - The tone should be professional, warm, and insightful — aimed at helping parents understand their child's daily growth.
        - REMEMBER: Always use "{user_info["student_name"]}" instead of any pronouns or names from the content

        Instructions for Report Generation
        Assign scores based on clear, evidence-backed observations for each area.

        Explain each score with a specific reason—avoid generalizations or repeated points. Every score must be justified individually and precisely.

        Use the following rating scale consistently:

        Ratings Scale:
         Excellent (7/7 areas) – Clear growth with strong evidence
         Good (5-6 areas) – Solid engagement with positive trends
         Fair (3-4 areas) – Some engagement, needs encouragement
         Needs Work (1-2 areas) – Area not activated or underperforming today
        Always include the complete legend in every report so the evaluator or reader can cross-check scores against the criteria.

        Ensure the entire report strictly follows the legend and that scoring aligns accurately with the defined scale.

        Do not use tables for the "Growth Metrics & Observations" section. Present the content in a well-spaced, structured paragraph format to preserve formatting integrity across platforms.
        📝 TEXT CONTENT:
        {text_content}

        🧾 Daily Insights Format for Parents

        🧒 Child's Name: {student_name}
        📅 Date: [{user_info.get("session_date", "Today")}]
        🌱 Curiosity Seed Explored: [Extract from text]

        📊 Growth Metrics & Observations
        Growth Area | Rating | Observation Summary
        🧠 Intellectual | [✅ Excellent/✅ Good/⚠️ Fair/📈 Needs Work] | [Brief summary]
        😊 Emotional | [✅ Excellent/✅ Good/⚠️ Fair/📈 Needs Work] | [Brief summary]
        🤝 Social | [✅ Excellent/✅ Good/⚠️ Fair/📈 Needs Work] | [Brief summary]
        🎨 Creativity | [✅ Excellent/✅ Good/⚠️ Fair/📈 Needs Work] | [Brief summary]
        🏃 Physical | [✅ Excellent/✅ Good/⚠️ Fair/📈 Needs Work] | [Brief summary]
        🚀 Planning/Independence | [✅ Excellent/✅ Good/⚠️ Fair/📈 Needs Work] | [Brief summary]
        🧭 Character | [✅ Excellent/✅ Good/⚠️ Fair/📈 Needs Work] | [Brief summary]
        
        🌈 Curiosity Response Index: [1-10] / 10  
        [Brief explanation of {student_name}'s engagement with the curiosity seed]ation: [Describe structure and coherence of thought process]  

        🧠 Overall Growth Score:  
        [🔵 Balanced Growth / 🟡 Moderate Growth / 🔴 Limited Growth] – [X/7] Areas Active 
        [Brief recommendation for next steps or continued development for {student_name}]

        📣 Note for Parent:  
        [Comprehensive summary for parents with actionable insights and encouragement based on today's session for {student_name}]

        🟢 Legend

        ✅ Performance by Area
        🟢 Excellent (7/7 areas) – Clear growth with strong evidence
        💚 Good (5-6 areas) – Solid engagement with positive trends        
        ⚠️ Fair (3-4 areas) – Some engagement, needs encouragement        
        📈 Needs Work (1-2 areas) – Area not activated or underperforming today

        give the entire report such that its a direct send worthy item, so all things should always be there and no other unecessary words in the response. No repetation.
        Also make sure each and every report generated always has the legend "🟢 Legend

        ✅ Performance by Area
        🟢 Excellent (7/7 areas) – Clear growth with strong evidence
        💚 Good (5-6 areas) – Solid engagement with positive trends        
        ⚠️ Fair (3-4 areas) – Some engagement, needs encouragement        
        📈 Needs Work (1-2 areas) – Area not activated or underperforming today" at the bottom of each report as the format of the report specifies. 
        """

        try:
            model = genai.GenerativeModel("gemini-2.0-flash")
            config = genai.types.GenerationConfig(temperature=0.2)
            response = model.generate_content(
                [{"role": "user", "parts": [{"text": prompt}]}],
                generation_config=config,
            )
            return response.text
        except Exception as e:
            return f"Error generating report: {str(e)}"

    def generate_ai_communication_review(self, transcript, user_info):
        """Generate AI communication review for peer review system"""
        # Get child information with gender
        from models.database import get_child_by_id

        child = get_child_by_id(user_info.get("child_id"))
        pronouns = (
            self.get_pronouns(child["gender"])
            if child and child.get("gender")
            else {"subject": "they", "object": "them", "possessive": "their"}
        )

        prompt = f"""
        You are an AI assistant analyzing observer-student communication sessions for educational quality assessment. 
        Generate a comprehensive communication review in a professional report format based on the provided transcript.

        STUDENT: {user_info["student_name"]}
        OBSERVER: {user_info["observer_name"]}
        TRANSCRIPT: {transcript}
        
        Use these pronouns for the student throughout the review: subject = {pronouns["subject"]}, object = {pronouns["object"]}, possessive = {pronouns["possessive"]}

        Generate a detailed AI communication review in the following professional format:

        Analysis of Observer's Communication Style
        This report analyzes the conversation transcripts to evaluate the observer's adherence to non-judgmental and non-teaching communication techniques with {user_info["student_name"]}.

        1. Instances of Direct Advice, Judgment, or Teaching
        [Analyze specific examples where the observer provided direct advice, made judgments, or acted as a teacher rather than a listener]

        ● 🚩 Red Flag: [Specific issue identified]
        ○ Instances: [Describe specific examples from the transcript]
        ■ [Date if available]: "[Exact quote from transcript]"
        ○ Analysis: [Explain why this is problematic and its impact]

        ● 🚩 Red Flag: [Another specific issue]
        ○ Instance: [Describe the instance]
        ○ Analysis: [Explain the impact and why it's concerning]

        2. Suggested Rephrasing for Non-Judgmental Communication
        [Provide specific examples of how the observer could rephrase their interactions]

        ● For [specific issue]:
        ○ Instead of: "[Current problematic phrase]"
        ○ Ask: "[Suggested non-judgmental alternative]"

        3. Missed Opportunities for Deeper Conversation
        [Identify moments where the observer could have explored topics more deeply]

        ● Date: [if available]
        ○ Missed Opportunity: [Describe what the student shared that could have been explored further]
        ○ Suggested Question: "[Specific question the observer could have asked]"

        4. Adherence to Non-Judgmental Listening: A Summary
        [Create a summary table of the observer's performance]

        Date | Adherence (✅/❌) | Red Flags | Suggested Improvements
        [Date] | ✅/❌ | Yes/No | [Brief description of what went well or needs improvement]

        5. Findings and Recommendations for Program Managers
        [Provide actionable insights and recommendations]

        Observer Improvement Areas
        [Summarize the main areas where the observer needs to improve]

        Impact on the Student ({user_info["student_name"]})
        [Analyze how the observer's communication style affects the student]

        Recommendation: [Provide specific, actionable recommendations for improvement]

        Format the review with proper hierarchical structure using the bullet points (●, ○, ■) and numbered sections as shown above. Include specific quotes from the transcript when relevant. Be constructive and educational in tone while being honest about areas that need improvement.
        """

        try:
            model = genai.GenerativeModel("gemini-2.0-flash")
            response = model.generate_content(
                [{"role": "user", "parts": [{"text": prompt}]}]
            )
            return response.text
        except Exception as e:
            return f"Error generating AI communication review: {str(e)}"

    def create_word_document_with_emojis(self, report_content):
        """Create a Word document from the report content with emoji support"""
        doc = docx.Document()

        # Set document encoding and font for emoji support
        style = doc.styles["Normal"]
        font = style.font
        font.name = "Segoe UI Emoji"

        title = doc.add_heading("📋 Daily Insights", 0)
        title.runs[0].font.name = "Segoe UI Emoji"

        # Process the report content line by line
        lines = report_content.split("\n")
        for line in lines:
            line = line.strip()
            if not line:
                continue

            # Clean up markdown formatting but preserve emojis
            line = line.replace("**", "")

            if line.startswith(("🧒", "📅", "🌱")):
                p = doc.add_paragraph()
                run = p.add_run(line)
                run.bold = True
                run.font.name = "Segoe UI Emoji"
                run.font.size = docx.shared.Pt(12)
            elif line.startswith("📊"):
                heading = doc.add_heading(line, level=1)
                heading.runs[0].font.name = "Segoe UI Emoji"
            elif line.startswith(("🧠", "😊", "🤝", "🎨", "🏃", "🧭", "🚀")):
                p = doc.add_paragraph()
                run = p.add_run(line)
                run.bold = True
                run.font.name = "Segoe UI Emoji"
                run.font.size = docx.shared.Pt(11)
            elif line.startswith(("🌈", "🗣️")):
                heading = doc.add_heading(line, level=2)
                heading.runs[0].font.name = "Segoe UI Emoji"
            elif line.startswith("🧠 Overall"):
                heading = doc.add_heading(line, level=2)
                heading.runs[0].font.name = "Segoe UI Emoji"
            elif line.startswith("📣"):
                heading = doc.add_heading(line, level=2)
                heading.runs[0].font.name = "Segoe UI Emoji"
            elif line.startswith("🟢 Legend"):
                heading = doc.add_heading(line, level=3)
                heading.runs[0].font.name = "Segoe UI Emoji"
            elif line.startswith(("✅", "⚠️", "📈", "🔵", "🟢", "🟡", "🔴", "•", "💚")):
                p = doc.add_paragraph(line, style="List Bullet")
                p.runs[0].font.name = "Segoe UI Emoji"
                p.runs[0].font.size = docx.shared.Pt(10)
            else:
                p = doc.add_paragraph()
                run = p.add_run(line)
                run.font.name = "Segoe UI Emoji"
                run.font.size = docx.shared.Pt(10)

        # Save to BytesIO object
        docx_bytes = io.BytesIO()
        doc.save(docx_bytes)
        docx_bytes.seek(0)

        return docx_bytes

    def create_pdf_alternative(self, report_content):
        """Create PDF using reportlab instead of WeasyPrint"""
        from reportlab.lib.pagesizes import letter, A4
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.lib.colors import black, blue, green, red

        # Create emoji mapping for text replacement
        emoji_map = {
            "🧒": "[Child]",
            "📅": "[Date]",
            "🌱": "[Curiosity Seed]",
            "📊": "[Growth Metrics]",
            "🧠": "[Intellectual]",
            "😊": "[Emotional]",
            "🤝": "[Social]",
            "🎨": "[Creative]",
            "🏃": "[Physical]",
            "🧭": "[Character/Values]",
            "🚀": "[Planning/Independence]",
            "🌈": "[Curiosity Response]",
            "🗣️": "[Communication Skills]",
            "📣": "[Note for Parent]",
            "🟢": "[Excellent]",
            "✅": "[Good]",
            "⚠️": "[Fair]",
            "📈": "[Needs Work]",
            "🔵": "[Balanced Growth]",
            "🟡": "[Moderate Growth]",
            "🔴": "[Limited Growth]",
            "💚": "[Good Score]",
            "📋": "[Report]",
        }

        # Replace emojis with readable text
        pdf_content = report_content
        for emoji, replacement in emoji_map.items():
            pdf_content = pdf_content.replace(emoji, replacement)

        buffer = io.BytesIO()
        doc = SimpleDocTemplate(
            buffer,
            pagesize=A4,
            rightMargin=72,
            leftMargin=72,
            topMargin=72,
            bottomMargin=18,
        )

        styles = getSampleStyleSheet()

        # Create custom styles
        title_style = ParagraphStyle(
            "CustomTitle",
            parent=styles["Heading1"],
            fontSize=16,
            spaceAfter=20,
            textColor=blue,
            alignment=1,  # Center alignment
        )

        heading_style = ParagraphStyle(
            "CustomHeading",
            parent=styles["Heading2"],
            fontSize=12,
            spaceAfter=10,
            textColor=black,
            fontName="Helvetica-Bold",
        )

        normal_style = ParagraphStyle(
            "CustomNormal",
            parent=styles["Normal"],
            fontSize=10,
            spaceAfter=4,
            fontName="Helvetica",
        )

        story = []

        # Add title
        title = Paragraph("[Report] Daily Insights", title_style)
        story.append(title)
        story.append(Spacer(1, 12))

        # Process report content
        lines = pdf_content.split("\n")
        for line in lines:
            line = line.strip()
            if not line:
                continue

            try:
                if line.startswith(("[Child]", "[Date]", "[Curiosity Seed]")):
                    para = Paragraph(f"<b>{line}</b>", heading_style)
                    story.append(para)
                elif line.startswith("[Growth Metrics]"):
                    para = Paragraph(f"<b>{line}</b>", heading_style)
                    story.append(para)
                elif line.startswith(
                    (
                        "[Intellectual]",
                        "[Emotional]",
                        "[Social]",
                        "[Creativity]",
                        "[Physical]",
                        "[Character/Values]",
                        "[Planning/Independence]",
                    )
                ):
                    para = Paragraph(f"<b>{line}</b>", normal_style)
                    story.append(para)
                elif line.startswith(
                    (
                        "[Curiosity Response]",
                        "[Communication Skills]",
                        "[Note for Parent]",
                    )
                ):
                    para = Paragraph(f"<b>{line}</b>", heading_style)
                    story.append(para)
                elif "Overall Growth Score" in line:
                    para = Paragraph(f"<b>{line}</b>", heading_style)
                    story.append(para)
                elif line.startswith("[Excellent] Legend"):
                    para = Paragraph(f"<b>{line}</b>", heading_style)
                    story.append(para)
                else:
                    para = Paragraph(line, normal_style)
                    story.append(para)

                story.append(Spacer(1, 4))

            except Exception as e:
                # Skip problematic lines
                continue

        doc.build(story)
        buffer.seek(0)

        return buffer

    def create_pdf_with_emojis(self, report_content):
        """Alias for create_pdf_alternative for backward compatibility"""
        return self.create_pdf_alternative(report_content)

    def create_word_document(self, report_content):
        """Legacy method - calls the emoji version"""
        return self.create_word_document_with_emojis(report_content)

    def send_email(self, recipient_email, subject, message):
        """Send email with the observation report"""
        sender_email = Config.EMAIL_USER
        sender_password = Config.EMAIL_PASSWORD

        if not sender_password:
            return False, "Email password not configured"

        smtp_server = "smtp.gmail.com"
        smtp_port = 587

        msg = MIMEMultipart()
        msg["From"] = sender_email
        msg["To"] = recipient_email
        msg["Subject"] = subject

        # Convert message to HTML format for better emoji display
        html_message = f"""
        <html>
        <head>
            <meta charset="UTF-8">
            <style>
                body {{ font-family: 'Segoe UI', Arial, sans-serif; line-height: 1.6; }}
                pre {{ white-space: pre-wrap; font-family: inherit; }}
            </style>
        </head>
        <body>
            <pre>{message}</pre>
        </body>
        </html>
        """

        msg.attach(MIMEText(html_message, "html", "utf-8"))

        try:
            server = smtplib.SMTP(smtp_server, smtp_port)
            server.starttls()
            server.login(sender_email, sender_password)
            server.send_message(msg)
            server.quit()
            return True, f"Email sent to {recipient_email}"
        except smtplib.SMTPAuthenticationError:
            return False, "Error: Authentication failed. Check your email and password."
        except smtplib.SMTPException as e:
            return False, f"Error: Failed to send email. {e}"
        except Exception as e:
            return False, f"Error: {str(e)}"

    def generate_custom_report_from_prompt(self, prompt, child_id):
        """Generate custom report based on prompt and stored data with new JSON format"""
        from models.database import get_supabase_client, get_child_by_id

        try:
            supabase = get_supabase_client()
            # Get all processed observations for the child
            processed_data = (
                supabase.table("processed_observations")
                .select("*")
                .eq("child_id", child_id)
                .execute()
            )
            observations = (
                supabase.table("observations")
                .select("*")
                .eq("student_id", child_id)
                .execute()
            )

            # Get child information with gender
            child = get_child_by_id(child_id)
            child_name = child["name"] if child else "Student"
            pronouns = (
                self.get_pronouns(child["gender"])
                if child and child.get("gender")
                else {"subject": "they", "object": "them", "possessive": "their"}
            )

            # Combine all data
            all_data = {
                "processed_observations": processed_data.data,
                "observations": observations.data,
            }

            custom_prompt = f"""
            You are an AI assistant for a learning observation system. Extract and structure information from the provided observation data based on the user's specific request.

CRITICAL INSTRUCTIONS FOR NAME AND GENDER USAGE:
- ALWAYS use the exact name from the database: {child_name}
- NEVER use any names that appear in the observation data or audio transcriptions
- Use these pronouns for the student throughout the report: subject = {pronouns["subject"]}, object = {pronouns["object"]}, possessive = {pronouns["possessive"]}
- When describing activities, use "{child_name}" or the appropriate pronouns ({pronouns["subject"]}/{pronouns["object"]}/{pronouns["possessive"]})

            Based on the following prompt and all available data for this child, generate a comprehensive custom report in the specified JSON format:

            USER PROMPT: {prompt}

            AVAILABLE DATA: {json.dumps(all_data, indent=2)}

            Format your response as JSON with the following structure:
            {{
              "studentName": "{child_name}",
              "studentId": "Custom Report ID",
              "className": "Custom Analysis Report",
              "date": "{datetime.now().strftime("%Y-%m-%d")}",
              "observations": "Detailed description combining all relevant observations that match the user's prompt - refer to '{child_name}' or 'the student'",
              "strengths": ["List of strengths observed in {child_name} based on available data"],
              "areasOfDevelopment": ["List of areas where {child_name} needs improvement"],
              "recommendations": ["List of recommended actions for {child_name} based on the prompt and data"]
            }}

            For observations, provide full detailed descriptions like:
            "{child_name} learned how to make maggi from their parent through in-person mode, including all steps from boiling water to adding spices"

            Be creative in extracting information based on context and ensure the response directly addresses the user's prompt: "{prompt}"

            REMEMBER: Always use "{child_name}" from the database, never names from observation data, and avoid gender assumptions.
            """

            model = genai.GenerativeModel("gemini-2.0-flash")
            response = model.generate_content(
                [{"role": "user", "parts": [{"text": custom_prompt}]}]
            )

            # Try to parse as JSON and format nicely
            try:
                # Clean the response text to extract JSON
                response_text = response.text.strip()

                # Remove markdown code blocks if present
                if response_text.startswith("```json"):
                    response_text = response_text[7:]  # Remove ```json
                if response_text.startswith("```"):
                    response_text = response_text[3:]  # Remove ```
                if response_text.endswith("```"):
                    response_text = response_text[:-3]  # Remove trailing ```

                # Find JSON object in the text
                start_idx = response_text.find("{")
                end_idx = response_text.rfind("}") + 1

                if start_idx != -1 and end_idx > start_idx:
                    json_text = response_text[start_idx:end_idx]
                    json_response = json.loads(json_text)
                else:
                    json_response = json.loads(response_text)

                # Format the JSON response into a readable report
                formatted_report = f"""
📋 Custom Report: {json_response.get("className", "Custom Analysis Report")}

🧒 Student Name: {json_response.get("studentName", child_name)}
📅 Date: {json_response.get("date", datetime.now().strftime("%Y-%m-%d"))}
📝 Report Type: Custom Analysis

📊 Observations Summary:
{json_response.get("observations", "No observations available")}

⭐ Strengths Identified:
{chr(10).join([f"• {strength}" for strength in json_response.get("strengths", [])])}

📈 Areas for Development:
{chr(10).join([f"• {area}" for area in json_response.get("areasOfDevelopment", [])])}

💡 Recommendations:
{chr(10).join([f"• {rec}" for rec in json_response.get("recommendations", [])])}

📋 Report Generated: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}
                """

                return formatted_report.strip()

            except (json.JSONDecodeError, ValueError) as e:
                # If JSON parsing fails, try to extract and format manually
                logger.error(f"JSON parsing failed: {str(e)}")
                logger.error(f"Response text: {response.text}")

                # Fallback: return a formatted version of the raw response
                return f"""
📋 Custom Report: Custom Analysis Report

🧒 Student Name: {child_name}
📅 Date: {datetime.now().strftime("%Y-%m-%d")}
📝 Report Type: Custom Analysis

📊 Analysis Results:
{response.text}

📋 Report Generated: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}
                """.strip()

        except Exception as e:
            return f"Error generating custom report: {str(e)}"

    def generate_monthly_report_json_format(
        self, observations, goal_progress, child_name, year, month, child_id=None
    ):
        """Generate monthly summary in the new JSON format with graph recommendations"""
        try:
            import calendar

            # Get child information with gender
            from models.database import get_child_by_id

            child = get_child_by_id(child_id) if child_id else None
            pronouns = (
                self.get_pronouns(child["gender"])
                if child and child.get("gender")
                else {"subject": "they", "object": "them", "possessive": "their"}
            )

            # Prepare data for analysis
            observation_texts = []
            all_strengths = []
            all_developments = []
            all_recommendations = []

            for obs in observations:
                observation_texts.append(obs.get("observations", ""))

                # Parse strengths, developments, and recommendations
                if obs.get("strengths"):
                    try:
                        strengths = (
                            json.loads(obs["strengths"])
                            if isinstance(obs["strengths"], str)
                            else obs["strengths"]
                        )
                        all_strengths.extend(strengths)
                    except:
                        pass

                if obs.get("areas_of_development"):
                    try:
                        developments = (
                            json.loads(obs["areas_of_development"])
                            if isinstance(obs["areas_of_development"], str)
                            else obs["areas_of_development"]
                        )
                        all_developments.extend(developments)
                    except:
                        pass

                if obs.get("recommendations"):
                    try:
                        recommendations = (
                            json.loads(obs["recommendations"])
                            if isinstance(obs["recommendations"], str)
                            else obs["recommendations"]
                        )
                        all_recommendations.extend(recommendations)
                    except:
                        pass

            # Calculate metrics for graphs
            total_observations = len(observations)
            active_goals = len(
                [g for g in goal_progress if g.get("status") == "active"]
            )
            completed_goals = len(
                [g for g in goal_progress if g.get("status") == "achieved"]
            )

            # Count frequency of strengths and development areas
            strength_counts = {}
            for strength in all_strengths:
                strength_counts[strength] = strength_counts.get(strength, 0) + 1

            development_counts = {}
            for dev in all_developments:
                development_counts[dev] = development_counts.get(dev, 0) + 1

            # Prepare graph data suggestions
            graph_suggestions = []

            if total_observations > 0:
                graph_suggestions.append(
                    {
                        "type": "bar_chart",
                        "title": "Observation Frequency by Week",
                        "description": f"Shows {total_observations} observations recorded throughout the month",
                    }
                )

            if strength_counts:
                graph_suggestions.append(
                    {
                        "type": "pie_chart",
                        "title": "Top Strengths Distribution",
                        "data": strength_counts,
                        "description": f"Distribution of {len(strength_counts)} different strength areas",
                    }
                )

            if development_counts:
                graph_suggestions.append(
                    {
                        "type": "horizontal_bar",
                        "title": "Development Areas Focus",
                        "data": development_counts,
                        "description": f"Areas requiring attention with frequency counts",
                    }
                )

            if goal_progress:
                graph_suggestions.append(
                    {
                        "type": "donut_chart",
                        "title": "Goal Progress Status",
                        "data": {"Active": active_goals, "Completed": completed_goals},
                        "description": f"Goal completion status: {completed_goals} completed, {active_goals} active",
                    }
                )

            # Create comprehensive prompt for JSON generation
            monthly_prompt = f"""
            You are an AI assistant for a learning observation system. Generate a comprehensive monthly report based on the provided observation data.

CRITICAL INSTRUCTIONS FOR NAME AND GENDER USAGE:
- ALWAYS use the exact name from the database: {child_name}
- NEVER use any names that appear in the observation data or audio transcriptions
- Use these pronouns for the student throughout the report: subject = {pronouns["subject"]}, object = {pronouns["object"]}, possessive = {pronouns["possessive"]}
- When describing activities and progress, use "{child_name}" or the appropriate pronouns ({pronouns["subject"]}/{pronouns["object"]}/{pronouns["possessive"]})

            MONTH: {calendar.month_name[month]} {year}
            STUDENT: {child_name}
            TOTAL OBSERVATIONS: {total_observations}
            GOALS STATUS: {active_goals} active, {completed_goals} completed

            OBSERVATION DATA: {json.dumps(observation_texts[:5], indent=2)}  # Limit for prompt size
            STRENGTHS IDENTIFIED: {list(strength_counts.keys())[:10]}
            DEVELOPMENT AREAS: {list(development_counts.keys())[:10]}

            QUANTIFIABLE METRICS FOR GRAPHS:
            {json.dumps(graph_suggestions, indent=2)}

            Format your response as JSON with the following structure:
            {{
              "studentName": "{child_name}",
              "studentId": "Monthly Report ID",
              "className": "Monthly Progress Summary",
              "date": "{calendar.month_name[month]} {year}",
              "observations": "Comprehensive monthly summary combining all {total_observations} observations, highlighting key learning moments for {child_name}, progress patterns, and notable developments throughout the month",
              "strengths": ["List of top strengths observed consistently in {child_name} throughout the month"],
              "areasOfDevelopment": ["List of areas where {child_name} needs continued focus and improvement"],
              "recommendations": ["List of specific recommended actions for {child_name} for the next month based on observed patterns"],
              "monthlyMetrics": {{
                "totalObservations": {total_observations},
                "activeGoals": {active_goals},
                "completedGoals": {completed_goals},
                "topStrengths": {dict(list(strength_counts.items())[:5])},
                "developmentFocus": {dict(list(development_counts.items())[:5])}
              }},
              "suggestedGraphs": {graph_suggestions}
            }}

            For observations, provide a comprehensive monthly summary like:
            "Throughout {calendar.month_name[month]}, {child_name} demonstrated consistent growth in multiple areas. Key learning highlights include [specific examples from observations]. {child_name} showed particular strength in [areas] while developing skills in [areas]. Notable progress was observed in [specific skills/subjects]."

            REMEMBER: Always use "{child_name}" from the database, never names from observation data, and avoid all gender assumptions.
            """

            # Generate the report using AI
            model = genai.GenerativeModel("gemini-2.0-flash")
            response = model.generate_content(
                [{"role": "user", "parts": [{"text": monthly_prompt}]}]
            )

            # Try to parse as JSON and format nicely
            try:
                json_response = json.loads(response.text)

                # Format the JSON response into a readable report
                formatted_report = f"""
📋 Monthly Report: {json_response.get("className", "Monthly Progress Summary")}

🧒 Student Name: {json_response.get("studentName", child_name)}
📅 Period: {json_response.get("date", f"{calendar.month_name[month]} {year}")}

📊 Monthly Metrics:
• Total Observations: {json_response.get("monthlyMetrics", {}).get("totalObservations", total_observations)}
• Active Goals: {json_response.get("monthlyMetrics", {}).get("activeGoals", active_goals)}
• Completed Goals: {json_response.get("monthlyMetrics", {}).get("completedGoals", completed_goals)}

📝 Monthly Observations Summary:
{json_response.get("observations", "No observations summary available")}

⭐ Strengths Observed:
{chr(10).join([f"• {strength}" for strength in json_response.get("strengths", [])])}

📈 Areas for Development:
{chr(10).join([f"• {area}" for area in json_response.get("areasOfDevelopment", [])])}

💡 Recommendations for Next Month:
{chr(10).join([f"• {rec}" for rec in json_response.get("recommendations", [])])}

📊 Suggested Visual Analytics:
{chr(10).join([f"• {graph['title']}: {graph['description']}" for graph in json_response.get("suggestedGraphs", [])])}

📋 Report Generated: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}
                """

                return formatted_report.strip()

            except json.JSONDecodeError:
                # If not valid JSON, return the raw response
                return response.text

        except Exception as e:
            return f"Error generating monthly summary: {str(e)}"

    def generate_monthly_docx_report(
        self,
        observations,
        goal_progress,
        strength_counts,
        development_counts,
        summary_json,
    ):
        """
        Generate a narrative-rich monthly report as a Word document, with embedded charts.
        """
        import docx
        from docx.shared import Inches, Pt
        from io import BytesIO
        import matplotlib.pyplot as plt
        from matplotlib.ticker import MaxNLocator
        import re
        import json

        doc = docx.Document()
        style = doc.styles["Normal"]
        font = style.font
        font.name = "Segoe UI"
        font.size = Pt(11)

        # --- Narrative Section ---
        doc.add_heading(f"Monthly Growth Report", 0)
        doc.add_paragraph(f"{summary_json.get('date', '')}")
        doc.add_paragraph(f"Student: {summary_json.get('studentName', '')}")
        doc.add_paragraph("")
        doc.add_paragraph(summary_json.get("observations", ""))
        doc.add_paragraph("")

        # --- Strengths ---
        doc.add_heading("Strengths Observed", level=1)
        for s in summary_json.get("strengths", []):
            doc.add_paragraph(s, style="List Bullet")
        doc.add_paragraph("")

        # --- Areas for Development ---
        doc.add_heading("Areas for Development", level=1)
        for a in summary_json.get("areasOfDevelopment", []):
            doc.add_paragraph(a, style="List Bullet")
        doc.add_paragraph("")

        # --- Recommendations ---
        doc.add_heading("Recommendations for Next Month", level=1)
        for r in summary_json.get("recommendations", []):
            doc.add_paragraph(r, style="List Bullet")
        doc.add_paragraph("")

        # ---  Analytics ---
        doc.add_heading("Learning Analytics", level=1)
        analytics = summary_json.get("learningAnalytics", {})
        for k, v in analytics.items():
            doc.add_paragraph(f"{k.replace('_', ' ').title()}: {v}")
        doc.add_paragraph("")

        # --- Progress Insights ---
        doc.add_heading("Progress Insights", level=1)
        for insight in summary_json.get("progressInsights", []):
            doc.add_paragraph(insight, style="List Bullet")
        doc.add_paragraph("")

        # --- Graphs Section ---
        doc.add_heading("Visual Analytics", level=1)

        # Curiosity and Growth scores by day (parse from observations)
        curiosity_by_date = {}
        growth_by_date = {}
        for obs in observations:
            date = obs.get("date")
            try:
                full_data = json.loads(obs.get("full_data", "{}"))
                report = full_data.get("formatted_report", "")
            except Exception:
                report = ""
            curiosity_match = re.search(
                r"🌈 Curiosity Response Index: (\\d{1,2}) ?/ ?10", report
            )
            if curiosity_match:
                curiosity_score = int(curiosity_match.group(1))
                curiosity_by_date[date] = curiosity_score
            growth_match = re.search(r"Overall Growth Score.*?(\\d)\\s*/\\s*7", report)
            if growth_match:
                growth_score = int(growth_match.group(1))
                growth_by_date[date] = growth_score
        # Sort by date
        curiosity_dates = sorted(curiosity_by_date.keys())
        growth_dates = sorted(growth_by_date.keys())
        curiosity_scores = [curiosity_by_date[d] for d in curiosity_dates]
        growth_scores = [growth_by_date[d] for d in growth_dates]

        # --- Curiosity Line Chart ---
        if curiosity_dates:
            fig, ax = plt.subplots()
            ax.plot(curiosity_dates, curiosity_scores, marker="o", color="blue")
            ax.set_title("Curiosity Response Index by Day")
            ax.set_xlabel("Date")
            ax.set_ylabel("Curiosity Score")
            ax.xaxis.set_major_locator(MaxNLocator(integer=True))
            plt.xticks(rotation=45, ha="right")
            plt.tight_layout()
            img_stream = BytesIO()
            plt.savefig(img_stream, format="png")
            plt.close(fig)
            img_stream.seek(0)
            doc.add_picture(img_stream, width=Inches(5.5))
            doc.add_paragraph("")

        # --- Growth Line Chart ---
        if growth_dates:
            fig, ax = plt.subplots()
            ax.plot(growth_dates, growth_scores, marker="o", color="green")
            ax.set_title("Overall Growth Score by Day")
            ax.set_xlabel("Date")
            ax.set_ylabel("Growth Score")
            ax.xaxis.set_major_locator(MaxNLocator(integer=True))
            plt.xticks(rotation=45, ha="right")
            plt.tight_layout()
            img_stream = BytesIO()
            plt.savefig(img_stream, format="png")
            plt.close(fig)
            img_stream.seek(0)
            doc.add_picture(img_stream, width=Inches(5.5))
            doc.add_paragraph("")

        # --- Other suggested graphs from summary_json ---
        for graph in summary_json.get("suggestedGraphs", []):
            if graph["type"] in ["line_chart", "bar_chart"]:
                fig, ax = plt.subplots()
                if graph["type"] == "line_chart":
                    x = list(graph["data"].keys())
                    y = list(graph["data"].values())
                    ax.plot(x, y, marker="o")
                elif graph["type"] == "bar_chart":
                    x = list(graph["data"].keys())
                    y = list(graph["data"].values())
                    ax.bar(x, y)
                ax.set_title(graph.get("title", ""))
                ax.set_xlabel(graph.get("xAxis", ""))
                ax.set_ylabel(graph.get("yAxis", ""))
                plt.xticks(rotation=45, ha="right")
                plt.tight_layout()
                img_stream = BytesIO()
                plt.savefig(img_stream, format="png")
                plt.close(fig)
                img_stream.seek(0)
                doc.add_picture(img_stream, width=Inches(5.5))
                doc.add_paragraph(graph.get("description", ""))
                doc.add_paragraph("")

        docx_bytes = BytesIO()
        doc.save(docx_bytes)
        docx_bytes.seek(0)
        return docx_bytes

    def generate_monthly_pdf_report(
        self,
        observations,
        goal_progress,
        strength_counts,
        development_counts,
        summary_json,
    ):
        """
        Generate a PDF version of the monthly report by converting the Word doc.
        """
        from docx2pdf import convert
        import tempfile
        import os
        from io import BytesIO

        try:
            # Generate the Word document first
            docx_bytes = self.generate_monthly_docx_report(
                observations,
                goal_progress,
                strength_counts,
                development_counts,
                summary_json,
            )

            # Create temporary files for conversion
            with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp_docx:
                tmp_docx.write(docx_bytes.read())
                tmp_docx_path = tmp_docx.name

            tmp_pdf_path = tmp_docx_path.replace(".docx", ".pdf")

            # Convert docx to pdf
            convert(tmp_docx_path, tmp_pdf_path)

            # Read the PDF file
            with open(tmp_pdf_path, "rb") as f:
                pdf_bytes = f.read()

            # Clean up temporary files
            try:
                os.remove(tmp_docx_path)
                os.remove(tmp_pdf_path)
            except OSError:
                pass  # Ignore cleanup errors

            return BytesIO(pdf_bytes)

        except Exception as e:
            # If PDF conversion fails, raise the error with more context
            raise Exception(
                f"PDF conversion failed: {str(e)}. This might be due to missing docx2pdf dependencies or system limitations."
            )

    def preprocess_audio_for_student(self, file, student_id):
        """Preprocess audio for students with known issues"""
        if student_id == "08cd0c39-62b1-4931-a9bb-1106a5206a39":
            # Add audio enhancement for this specific student
            # This could include noise reduction, volume normalization, etc.
            logger.info(
                "Applying audio preprocessing for student with detection issues"
            )
            # Return processed file or original if preprocessing fails
            pass
        return file

    def transcribe_with_whisper_fallback(self, file):
        """Fallback transcription method using Whisper or similar"""
        # Implement a backup transcription service
        # This is a placeholder - you'll need to implement actual fallback logic
        raise NotImplementedError("Fallback transcription not implemented yet")

    def generate_topic_suggestions(
        self, observer_data, child_data, child_name, child_id=None
    ):
        """Generate topic suggestions using Gemini AI based on observation history"""
        try:
            # Get child information with gender
            from models.database import get_child_by_id

            child = get_child_by_id(child_id) if child_id else None
            pronouns = (
                self.get_pronouns(child["gender"])
                if child and child.get("gender")
                else {"subject": "they", "object": "them", "possessive": "their"}
            )

            # Prepare data for analysis
            recent_themes = []
            recent_strengths = []
            recent_developments = []
            recent_curiosities = []
            learning_patterns = []

            # Process observer's recent observations
            for obs in observer_data:
                if obs.get("theme_of_day"):
                    recent_themes.append(obs["theme_of_day"])
                if obs.get("curiosity_seed"):
                    recent_curiosities.append(obs["curiosity_seed"])

                # Parse JSON fields safely
                try:
                    if obs.get("strengths"):
                        strengths = (
                            json.loads(obs["strengths"])
                            if isinstance(obs["strengths"], str)
                            else obs["strengths"]
                        )
                        recent_strengths.extend(
                            strengths[:3]
                        )  # Top 3 strengths per observation

                    if obs.get("areas_of_development"):
                        developments = (
                            json.loads(obs["areas_of_development"])
                            if isinstance(obs["areas_of_development"], str)
                            else obs["areas_of_development"]
                        )
                        recent_developments.extend(
                            developments[:2]
                        )  # Top 2 development areas
                except:
                    continue

                # Extract learning patterns from observations
                if obs.get("observations"):
                    learning_patterns.append(obs["observations"][:200] + "...")

            # Process child-specific data
            child_themes = []
            child_curiosities = []
            child_patterns = []

            for obs in child_data:
                if obs.get("theme_of_day"):
                    child_themes.append(obs["theme_of_day"])
                if obs.get("curiosity_seed"):
                    child_curiosities.append(obs["curiosity_seed"])
                if obs.get("observations"):
                    child_patterns.append(obs["observations"][:150] + "...")

            # Create comprehensive prompt for Gemini
            prompt = f"""
You are an educational consultant helping an Observer plan engaging learning sessions. Based on the learning history below, suggest 5-7 specific, actionable topics or activities for today's observation session with {child_name}.
Also, provide a brief rationale for each suggestion. For example, indicate which theme, curiosity, past mention/instance, or learning pattern from the child's or Observer's history informed your recommendation.

Use these pronouns for the student throughout the suggestions: subject = {pronouns["subject"]}, object = {pronouns["object"]}, possessive = {pronouns["possessive"]}

OBSERVER'S RECENT TEACHING HISTORY:
Recent Themes Covered: {", ".join(recent_themes[-8:]) if recent_themes else "None available"}
Recent Curiosity Seeds: {", ".join(recent_curiosities[-5:]) if recent_curiosities else "None available"}
Observed Strengths: {", ".join(list(set(recent_strengths[-10:]))) if recent_strengths else "None available"}
Areas for Development: {", ".join(list(set(recent_developments[-8:]))) if recent_developments else "None available"}

CHILD'S SPECIFIC LEARNING HISTORY:
Child's Previous Themes: {", ".join(child_themes[-6:]) if child_themes else "None available"}
Child's Curiosity Patterns: {", ".join(child_curiosities[-4:]) if child_curiosities else "None available"}
Recent Learning Patterns: {chr(10).join(child_patterns[-3:]) if child_patterns else "None available"}

GUIDELINES FOR SUGGESTIONS:
1. Build upon previous themes but introduce fresh perspectives
2. Address identified development areas through engaging activities
3. Leverage the child's demonstrated strengths and interests
4. Suggest age-appropriate, hands-on learning experiences
5. Include variety: academic, creative, social-emotional, and practical life skills
6. Consider seasonal relevance and current events when appropriate
7. Ensure topics can be explored in a 30-45 minute session
8. IMPORTANT: The conversation between teacher and student is happening over a voice call, so suggest activities that are possible and engaging in a remote setting
9. Also, provide proper reasoning or a reference for why you are suggesting the particular activity based on the past reports (the activity should have some co-relation to the past reports).

Please provide suggestions in this exact format:

🎯 **SUGGESTED TOPICS FOR TODAY'S SESSION**

1. **[Topic Title]** - [Brief description of the activity and learning objective]
2. **[Topic Title]** - [Brief description of the activity and learning objective]
3. **[Topic Title]** - [Brief description of the activity and learning objective]
4. **[Topic Title]** - [Brief description of the activity and learning objective]
5. **[Topic Title]** - [Brief description of the activity and learning objective]

💡 **FOCUS AREAS TO EMPHASIZE:**
- [Specific strength to build upon]
- [Development area to address]
- [Curiosity pattern to explore further]

🌟 **SESSION TIP:** [One practical tip for making today's session particularly engaging based on the child's learning patterns]
"""

            # Generate suggestions using Gemini
            model = genai.GenerativeModel("gemini-2.0-flash")
            response = model.generate_content(
                [{"role": "user", "parts": [{"text": prompt}]}]
            )

            if response and response.text:
                return response.text.strip()
            else:
                return self._fallback_suggestions(child_name)

        except Exception as e:
            logger.error(f"Error generating topic suggestions: {str(e)}")
            return self._fallback_suggestions(child_name)

    def _fallback_suggestions(self, child_name, child_id=None):
        """Fallback suggestions when AI fails"""
        # Get child information with gender
        from models.database import get_child_by_id

        child = get_child_by_id(child_id) if child_id else None
        pronouns = (
            self.get_pronouns(child["gender"])
            if child and child.get("gender")
            else {"subject": "they", "object": "them", "possessive": "their"}
        )

        return f"""🎯 **SUGGESTED TOPICS FOR TODAY'S SESSION**

1. **Creative Storytelling** - Have {child_name} create and narrate a story using everyday objects
2. **Nature Exploration** - Observe and discuss plants, weather, or seasonal changes
3. **Math in Daily Life** - Practice counting, sorting, or measuring with household items
4. **Science Experiments** - Simple experiments using safe household materials
5. **Cultural Learning** - Explore traditions, festivals, or geography through interactive discussion

💡 **FOCUS AREAS TO EMPHASIZE:**
- Encourage curiosity and questioning
- Build confidence through hands-on activities
- Develop communication and expression skills

🌟 **SESSION TIP:** Start with what interests {child_name} most and build the lesson around {pronouns["possessive"]} natural curiosity!"""
